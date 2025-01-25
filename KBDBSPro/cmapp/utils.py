from .models import currencystate
from .models import templist
import pandas as pd
from openpyxl import Workbook
from datetime import datetime
from django.db import models
from django.shortcuts import get_object_or_404
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from django.http import HttpResponse
from io import BytesIO
import base64


def get_currency_state():   
    currency_state_obj = currencystate.objects.first()

    if not currency_state_obj:
        return {}  # Return an empty dictionary if no record exists

    # Mapping fields from the model to a dictionary
    currency_state = {
        'fivehundred': currency_state_obj.fivehundred,
        'twohundred': currency_state_obj.twohundred,
        'onehundred': currency_state_obj.onehundred,
        'fifty': currency_state_obj.fifty,
        'twenty': currency_state_obj.twenty,
        'ten': currency_state_obj.ten,
        'five': currency_state_obj.five,
        'two': currency_state_obj.two,
        'one': currency_state_obj.one,
    }
    return currency_state  # Return the dictionary, not the model class


def display_bal():
    currency_state = get_currency_state()

    if not currency_state:
        return {"message": "No currency state available.", "total": 0}

    # Mapping from string denominations to numeric values
    denomination_mapping = {
        'fivehundred': 500,
        'twohundred': 200,
        'onehundred': 100,
        'fifty': 50,
        'twenty': 20,
        'ten': 10,
        'five': 5,
        'two': 2,
        'one': 1
    }

    # Initialize results
    denomination_details = {}
    total_balance = 0
    total_notes = 0

    # Calculate the total balance, breakdown, and total notes
    for denomination, count in currency_state.items():
        # Ensure the count is an integer, defaulting to 0 if it's None
        count = count or 0

        # Convert denomination to numeric value using the mapping
        numeric_denomination = denomination_mapping.get(denomination)
        if numeric_denomination:
            # Calculate the total for each denomination
            total_for_denomination = numeric_denomination * count
            denomination_details[numeric_denomination] = {
                "count": count,
                "total": total_for_denomination
            }

            # Add to the overall totals
            total_balance += total_for_denomination
            total_notes += count

    # Prepare the result
    bal_result = {
        "denomination_details": denomination_details,
        "total_balance": total_balance,
        "total_notes": total_notes
    }
    return bal_result

#---------Excel TO DB------------


def excel_to_db(uploaded_file):
    templist.objects.all().delete()
    # Read the uploaded Excel file
    df = pd.read_excel(uploaded_file)

    df.columns = ['Date', 'Name', 'Department_Name', 'Month', 'Amount', 'Note']
    df['Date'] = pd.to_datetime(df['Date'], errors='coerce').dt.date
    df['Amount'] = df['Amount'].fillna(0).astype(int)

    # Insert data into the database
    for _, row in df.iterrows():
        templist.objects.create(
            date=row['Date'],
            name=row['Name'],
            department_name=row['Department_Name'],
            month=row['Month'],
            amount=row['Amount'],
            note=row['Note']
        )
    print("excel to db completed")
    results = get_change_distribution()

    excel_file, excel_path = export_to_excel(results)

    return excel_file, excel_path
    

            
def calculate_change(amount):
    # Fetch the currency state from the database (id=1)
    currency_state_obj = get_object_or_404(currencystate, id=1)

    # Create a dictionary to map denominations to their counts
    currency_state = {
        500: currency_state_obj.fivehundred,
        200: currency_state_obj.twohundred,
        100: currency_state_obj.onehundred,
        50: currency_state_obj.fifty,
        20: currency_state_obj.twenty,
        10: currency_state_obj.ten,
        5: currency_state_obj.five,
        2: currency_state_obj.two,
        1: currency_state_obj.one,
    }

    # Initialize the distribution of change
    change_distribution = {denomination: 0 for denomination in currency_state.keys()}

    # Calculate change
    for denomination, count in sorted(currency_state.items(), key=lambda x: -x[0]):
        if amount <= 0:
            break
        num_notes = min(amount // denomination, count)
        if num_notes > 0:
            change_distribution[denomination] = num_notes
            amount -= denomination * num_notes
            currency_state[denomination] -= num_notes  # Update in memory

    # Check if exact change was possible
    if amount > 0:
        return "Insufficient currency to provide exact change."

    # Update the database with the new currency state
    currencystate.objects.filter(id=1).update(
        fivehundred=currency_state[500],
        twohundred=currency_state[200],
        onehundred=currency_state[100],
        fifty=currency_state[50],
        twenty=currency_state[20],
        ten=currency_state[10],
        five=currency_state[5],
        two=currency_state[2],
        one=currency_state[1],
    )

    # Remove unused denominations from the result
    change_distribution = {denom: count for denom, count in change_distribution.items() if count > 0}

    return change_distribution



def get_change_distribution():
    try:
        print("Try block of get_change_distribution..")
        
        # Fetch the current currency state
        currency_state_obj = currencystate.objects.first()
        print(f"Currency state object: {currency_state_obj}")  # Check if it's None or an object
        
        if not currency_state_obj:
            return {}  # Return an empty dictionary if no record exists

        # Mapping fields from the model to a dictionary
        currency_state = {
            'fivehundred': currency_state_obj.fivehundred,
            'twohundred': currency_state_obj.twohundred,
            'onehundred': currency_state_obj.onehundred,
            'fifty': currency_state_obj.fifty,
            'twenty': currency_state_obj.twenty,
            'ten': currency_state_obj.ten,
            'five': currency_state_obj.five,
            'two': currency_state_obj.two,
            'one': currency_state_obj.one,
        }
        
        print(f"Currency state dictionary: {currency_state}")

        # Calculate available balance
        available_balance = (
            currency_state['fivehundred'] * 500 +
            currency_state['twohundred'] * 200 +
            currency_state['onehundred'] * 100 +
            currency_state['fifty'] * 50 +
            currency_state['twenty'] * 20 +
            currency_state['ten'] * 10 +
            currency_state['five'] * 5 +
            currency_state['two'] * 2 +
            currency_state['one'] * 1
        )
        print(f"Available balance: {available_balance}")

        # Calculate total required amount
        total_required_amount = templist.objects.aggregate(total=models.Sum('amount'))['total']
        print(f"Total required amount: {total_required_amount}")

        if available_balance < total_required_amount:
            return "Insufficient balance to produce change"
        
        # Generate change for each entry in TempList
        results = []
        for entry in templist.objects.all():
            print(f"Processing entry: {entry}")  # Ensure records are being processed
            change = calculate_change(entry.amount)
            results.append({
                "date": entry.date,
                "name": entry.name,
                "department_name": entry.department_name,
                "month": entry.month,
                "amount": entry.amount,
                "note": entry.note,
                "change": change if isinstance(change, dict) else None
            })
        
        # Export to Excel
        print("Exporting to Excel..")
        export_to_excel(results)

        return results

    except Exception as e:
        print(f"Error in get_change_distribution: {e}")
        return []

#-----------Currency Convert-------------
marathi_units = ["", "एक", "दोन", "तीन", "चार", "पाच", "सहा", "सात", "आठ", "नऊ"]
marathi_tens = ["", "दहा", "वीस", "तीस", "चाळीस", "पन्नास", "साठ", "सत्तर", "ऐंशी", "नव्वद"]
marathi_teens = ["", "अकरा", "बारा", "तेरा", "चौदा", "पंधरा", "सोळा", "सतरा", "अठरा", "एकोणीस"]
marathi_tens_ext = {
    21: "एकवीस", 22: "बावीस", 23: "तेवीस", 24: "चोवीस", 25: "पंचवीस",
    26: "सव्वीस", 27: "सत्तावीस", 28: "अठ्ठावीस", 29: "एकोणतीस",
    31: "एकतीस", 32: "बत्तीस", 33: "तेहतीस", 34: "चौतीस", 35: "पस्तीस",
    36: "छत्तीस", 37: "सदोतीस", 38: "अडतीस", 39: "एकोणचाळीस",
    41: "एक्केचाळीस", 42: "बेचाळीस", 43: "त्रेचाळीस", 44: "चव्वेचाळीस",
    45: "पंचेचाळीस", 46: "शेहचाळीस", 47: "सत्तेचाळीस", 48: "अठ्ठेचाळीस", 49: "एकोणपन्नास",
    51: "एकावन्न", 52: "बाव्वन्न", 53: "त्रेपन्न", 54: "चौपन्न", 55: "पंचावन्न",
    56: "छप्पन", 57: "सत्तावन्न", 58: "अठ्ठावन्न", 59: "एकोणसाठ",
    61: "एकसष्ट", 62: "बासष्ट", 63: "तेसष्ट", 64: "चौसष्ट", 65: "पासष्ट",
    66: "सहासष्ठ", 67: "सदुसष्ट", 68: "अडुसष्ट", 69: "एकोणसत्तर",
    71: "एकाहत्तर", 72: "बहात्तर", 73: "त्र्याहत्तर", 74: "चौर्‍याहत्तर", 75: "पंचाहत्तर",
    76: "शहात्तर", 77: "सत्याहत्तर", 78: "अठ्ठ्याहत्तर", 79: "एकोणऐंशी",
    81: "एक्याऐंशी", 82: "ब्याऐंशी", 83: "त्र्याऐंशी", 84: "चौरेऐंशी", 85: "पंच्याऐंशी",
    86: "शहाऐंशी", 87: "सत्याऐंशी", 88: "अठ्ठ्याऐंशी", 89: "एकोणनव्वद",
    91: "एक्याण्णव", 92: "ब्याण्णव", 93: "त्र्याण्णव", 94: "चौऱ्याण्णव", 95: "पंच्याण्णव",
    96: "शहाण्णव", 97: "सत्त्याण्णव", 98: "अठ्ठ्याण्णव", 99: "नव्व्याण्णव"
}
marathi_hundreds = ["", "शंभर", "दोनशे", "तीनशे", "चारशे", "पाचशे", "सहाशे", "सातशे", "आठशे", "नऊशे"]
marathi_thousands = ["", "हजार", "लाख", "कोटी"]

# Updated function to convert numbers to Marathi
def number_to_marathi(num):
    if num == 0:
        return "शून्य"
    
    result = []

    # Handle crores
    if num >= 10000000:
        crores = num // 10000000
        if crores > 0:
            result.append(number_to_marathi(crores) + " कोटी")
        num %= 10000000

    # Handle millions (lakhs)
    if num >= 100000:
        lakhs = num // 100000
        if lakhs == 1:
            result.append("एक लाख")
        else:
            result.append(number_to_marathi(lakhs) + " लाख")
        num %= 100000

    # Handle thousands
    if num >= 1000:
        thousands = num // 1000
        result.append(number_to_marathi(thousands) + " हजार")
        num %= 1000

    # Handle hundreds
    if num >= 100:
        hundreds = num // 100
        result.append(marathi_hundreds[hundreds])
        num %= 100

    # Handle special case for numbers in tens (10–99) or teen numbers (11–19)
    if 10 < num < 20:
        result.append(marathi_teens[num - 10])
        num = 0
    elif num >= 10 and num < 100:
        tens = num // 10
        if num in marathi_tens_ext:
            result.append(marathi_tens_ext[num])
            num = 0
        else:
            result.append(marathi_tens[tens])
            num %= 10

    # Handle units (0–9)
    if num > 0:
        result.append(marathi_units[num])

    return " ".join(result) 


#-----------------------------------------------
def export_to_excel(results):
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Change Distribution"
    
    # Headers
    headers = ["Date", "Name", "Department Name", "Month", "Amount", "Note", "500", "200", "100", "50", "20", "10", "5", "2", "1", "Change Total", "C_amount", "Stamp", "Match"]
    sheet.append(headers)

    # Initialize counters for totals
    total_amount = 0
    total_change_counts = {500: 0, 200: 0, 100: 0, 50: 0, 20: 0, 10: 0, 5: 0, 2: 0, 1: 0}
    total_change_value = 0

    # Populate rows
    for result in results:
        date = result["date"]
        name = result["name"]
        department_name = result["department_name"]
        month = result["month"]
        amount = result["amount"]
        note = result["note"]
        change = result["change"]
        
        row = [date, name, department_name, month, amount, note]
        change_total = 0  # Initialize the change total for this person
        
        # Add to the overall amount total
        total_amount += amount

        # Check if change is available
        if change:
            for denomination in [500, 200, 100, 50, 20, 10, 5, 2, 1]:
                count = change.get(denomination, 0)
                row.append(count)
                change_total += denomination * count
                
                # Add to the overall denomination totals
                total_change_counts[denomination] += count
            total_change_value += change_total
            # Convert the change total to words and add to the C_amount column
            c_amount = number_to_marathi(change_total) + " रुपये"
        else:
            row.extend(["Insufficient"] * 9)
            c_amount = "Insufficient"
        
        row.append(change_total if change else "Insufficient")
        row.append(c_amount)
        
        # Add "Stamp" column based on amount
        if amount >= 5000:
            row.append(1)
        else:
            row.append(0)
        
        # Add "Match" column based on comparison between amount and change_total
        if amount == change_total:
            row.append("-")
        else:
            row.append("Wrong")

        sheet.append(row)

    # Add a final row for the totals
    total_row = ["Total", "", "", "", total_amount, ""]
    for denomination in [500, 200, 100, 50, 20, 10, 5, 2, 1]:
        total_row.append(total_change_counts[denomination])
    total_row.append(total_change_value)
    total_row.append("")  # Leave C_amount blank for the total row
    total_row.append("")  # Leave "Stamp" blank for the total row
    total_row.append("")  # Leave "Match" blank for the total row

    # Append the total row to the sheet
    sheet.append(total_row)
    
    # Generate the file name with the current date and time
    current_time = datetime.now().strftime("%d-%m-%y %I%M%p")
    excel_path = f"ChangeDistribution_{current_time}.xlsx"

    # Save the workbook to an in-memory file
    
    excel_file = BytesIO()
    workbook.save(excel_file)
    excel_file.seek(0)
    return excel_file, excel_path




def CreateVoucher(excel_file_bytes):
    try:
        # Load the Excel data directly from the binary bytes
        df = pd.read_excel(BytesIO(excel_file_bytes))
        df['Date'] = pd.to_datetime(df['Date'], errors='coerce')  # Convert to datetime, set invalid dates to NaT

        # Create a new Word document
        doc = Document()

        # Set A4 page size and narrow margins
        sections = doc.sections
        for section in sections:
            section.page_height = Inches(11.7)  # A4 page height in inches
            section.page_width = Inches(8.3)    # A4 page width in inches
            section.top_margin = Inches(0.5)    # Narrow top margin
            section.bottom_margin = Inches(0.1) # Narrow bottom margin
            section.left_margin = Inches(0.5)   # Narrow left margin
            section.right_margin = Inches(0.5)  # Narrow right margin

        # Function to set font to Mangal for Marathi text and bold headings
        def set_font(paragraph, font_name, font_size, bold=False):
            run = paragraph.runs[0]
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold
            r = run._r
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        # Function to safely format dates
        def format_date(date):
            if pd.isna(date):
                return "-"  # Or some other default value
            return date.strftime("%d/%m/%Y")

        # Add content to the document with a single border for each voucher
        voucher_count = 0
        for index, row in df.iterrows():
            if voucher_count == 2:
                doc.add_page_break()  # Add a page break after two vouchers
                voucher_count = 0

            # Create a table with one row and one column to contain the voucher text
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'  # Apply a grid style to add a border around the table
            cell = table.cell(0, 0)

            # Add voucher title
            paragraph = cell.add_paragraph("व्हाऊचर")
            set_font(paragraph, "Mangal", 14, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the title

            # Add "व्हाऊचर नं." on the left and "तारीख" on the right
            voucher_table = cell.add_table(rows=1, cols=2)
            voucher_table.autofit = True

            # Format the date
            date_only = format_date(row['Date'])

            # Left cell for "व्हाऊचर नं."
            voucher_left = voucher_table.cell(0, 0).paragraphs[0]
            voucher_left.add_run("व्हाऊचर नं. ").bold = True
            set_font(voucher_left, "Mangal", 12)

            # Right cell for "तारीख"
            voucher_right = voucher_table.cell(0, 1).paragraphs[0]
            voucher_right.add_run("तारीख: ").bold = True
            voucher_right.add_run(f"{date_only}")
            set_font(voucher_right, "Mangal", 12)
            voucher_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Add additional voucher details
            paragraph = cell.add_paragraph("व्हाऊचर लिहुन घेणार - ")
            set_font(paragraph, "Mangal", 12, bold=True)
            paragraph.add_run("कृष्णा बिझनेस डेव्हलपमेंट प्रा.लि.")

            paragraph = cell.add_paragraph("व्हाऊचर लिहुन देणार - ")
            set_font(paragraph, "Mangal", 12, bold=True)
            paragraph.add_run(f"{row['Name']}")

            paragraph = cell.add_paragraph("खात्याचे नाव - ")
            set_font(paragraph, "Mangal", 11)
            paragraph.add_run(f"Salary A/C ").bold = True
            paragraph.add_run(f"                        Department:  ")
            paragraph.add_run(f"{row['Department Name']}").bold = True

            paragraph = cell.add_paragraph("व्हाऊचर लिहुन देतो की, आपणाकडुन खालील तपशिलाप्रमाणे रक्कम मिळाली.")
            set_font(paragraph, "Mangal", 12)

            paragraph = cell.add_paragraph("तपशील:")
            set_font(paragraph, "Mangal", 12, bold=True)

            paragraph = cell.add_paragraph("Being the salary for the month ")
            set_font(paragraph, "Mangal", 12)
            paragraph.add_run(f"{row['Month']}").bold = True
            paragraph.add_run("-25").bold = True
            paragraph.add_run(" given to ")
            paragraph.add_run(f"{row['Name']}").bold = True
            paragraph2 = cell.add_paragraph("नोट: ")
            set_font(paragraph2, "Mangal", 9,)
            paragraph2.add_run(f"{row['Note']}__")
            set_font(paragraph2, "Mangal", 9)

            paragraph = cell.add_paragraph("रुपये:  ")
            set_font(paragraph, "Mangal", 12)
            paragraph.add_run(f"{row['Amount']} /-").bold = True
            set_font(paragraph, "Mangal", 12)

            paragraph = cell.add_paragraph("अक्षरी रुपये: ")
            set_font(paragraph, "Mangal", 12)
            paragraph.add_run(f"{row['C_amount']}").bold = True
            
            # Add the table for the denomination details
            denom_table = cell.add_table(rows=2, cols=10)  # 1 row for headers, 1 row for data
            denom_table.style = 'Table Grid'

            # Header row (500 200 100 50 20 10 5 2 1 Change Total)
            header_cells = denom_table.rows[0].cells
            header_cells[0].text = "500"
            header_cells[1].text = "200"
            header_cells[2].text = "100"
            header_cells[3].text = "50"
            header_cells[4].text = "20"
            header_cells[5].text = "10"
            header_cells[6].text = "5"
            header_cells[7].text = "2"
            header_cells[8].text = "1"
            header_cells[9].text = "Total"

            # Data row (values from the new columns)
            data_cells = denom_table.rows[1].cells
            data_cells[0].text = str(row['500'])
            data_cells[1].text = str(row['200'])
            data_cells[2].text = str(row['100'])
            data_cells[3].text = str(row['50'])
            data_cells[4].text = str(row['20'])
            data_cells[5].text = str(row['10'])
            data_cells[6].text = str(row['5'])
            data_cells[7].text = str(row['2'])
            data_cells[8].text = str(row['1'])
            data_cells[9].text = str(row['Change Total'])

            sig_table = cell.add_table(rows=2, cols=2)
            sig_table.autofit = True
            # Check the "Stamp" column value
            if row['Stamp'] == 1:
                # Create a new table within the cell for the checkmark box
                sig_right = sig_table.cell(0, 1).add_table(rows=1, cols=1)
                sig_right.style = 'Table Grid'  # Keep grid style but adjust the size
                box_cell = sig_right.cell(0, 0)
                
                # Set the cell content as a checkmark
                box_cell.text = "Stamp"
                
                # Set the width and height of the cell to make it a small box
                box_cell.width = Inches(0.5)  # Adjust the width as needed
                box_cell.height = Inches(0.5)  # Adjust the height as needed
                sig_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Center the checkmark in the cell
                box_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Set font size and bold for the checkmark
                set_font(box_cell.paragraphs[0], "Mangal", 8)
            else:
                # Create a new table within the cell for the checkmark box
                sig_right = sig_table.cell(0, 1).add_table(rows=1, cols=1)        
                box_cell = sig_right.cell(0, 0)
                
                # Set the cell content as a checkmark
                box_cell.text = "."
                
                # Set the width and height of the cell to make it a small box
                box_cell.width = Inches(0.5)  # Adjust the width as needed
                box_cell.height = Inches(0.5)  # Adjust the height as needed
                sig_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Center the checkmark in the cell
                box_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Set font size and bold for the checkmark
                set_font(box_cell.paragraphs[0], "Mangal", 3)

            # Left cell for "चेक्ड बाय"
            sig_left = sig_table.cell(1, 0).paragraphs[0]
            sig_left.add_run("चेक्ड बाय").bold = True
            set_font(sig_left, "Mangal", 12)
            sig_left.add_run(f"{row['Match']}").bold = True
            sig_left.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Right cell for "पैसे घेणाऱ्याची सही"
            sig_right_text = sig_table.cell(1, 1).paragraphs[0]
            sig_right_text.add_run("पैसे घेणाऱ्याची सही").bold = True
            set_font(sig_right_text, "Mangal", 12)
            sig_right_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Increment voucher count to track two vouchers per page
            voucher_count += 1

            # Add a blank paragraph between vouchers if it's the first on the page
            if voucher_count == 1:
                doc.add_paragraph()       
                

        # Save the final document
        current_time = datetime.now().strftime("%d-%m-%y %I%M%p")
        docx_filename = f"Voucher_Report_{current_time}.docx"
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        return doc_io.read(), docx_filename
    except Exception as e:
        print(f"Error generating document: {e}")
        return None, None